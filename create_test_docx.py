#!/usr/bin/env python3
"""Create a test DOCX file for RAG testing"""

from docx import Document
from pathlib import Path

def create_test_docx():
    doc = Document()
    
    # Add title
    title = doc.add_heading('RAG System Test Document', 0)
    
    # Add introduction
    intro = doc.add_paragraph(
        'This is a test Microsoft Word document created to demonstrate the RAG system\'s '
        'ability to process and index DOCX files. The document contains various sections '
        'that should be searchable through the RAG system.'
    )
    
    # Add a section header
    doc.add_heading('Technical Specifications', level=1)
    
    # Add technical content
    doc.add_paragraph(
        'The RAG (Retrieval-Augmented Generation) system uses sentence transformers '
        'to create embeddings of document content. These embeddings enable semantic '
        'search capabilities that go beyond simple keyword matching.'
    )
    
    # Add a list
    doc.add_paragraph(
        'Key features of the RAG system include:'
    )
    
    features = [
        'Document type support: DOCX, PDF, and text files',
        'Automatic text extraction with metadata preservation',
        'Smart chunking for large documents',
        'Semantic search using embeddings',
        'Fallback keyword search',
        'Source attribution in responses'
    ]
    
    for feature in features:
        p = doc.add_paragraph(feature, style='List Bullet')
    
    # Add another section
    doc.add_heading('Implementation Details', level=1)
    
    doc.add_paragraph(
        'The system is implemented using Python with the following key libraries:'
    )
    
    # Add a table
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Library'
    hdr_cells[1].text = 'Purpose'
    
    libraries = [
        ('sentence-transformers', 'Text embeddings and semantic search'),
        ('python-docx', 'Microsoft Word document processing'),
        ('PyPDF2', 'PDF document text extraction'),
        ('SQLite3', 'Knowledge base storage'),
        ('numpy', 'Numerical operations for embeddings')
    ]
    
    for lib, purpose in libraries:
        row_cells = table.add_row().cells
        row_cells[0].text = lib
        row_cells[1].text = purpose
    
    # Add conclusion
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(
        'This test document demonstrates the RAG system\'s ability to extract and '
        'index content from Microsoft Word documents, including text, tables, and '
        'structured information. When indexed, this content should be retrievable '
        'through semantic queries about RAG functionality, document processing, '
        'or implementation details.'
    )
    
    # Save the document
    output_path = Path('test_documents/rag_test_document.docx')
    doc.save(str(output_path))
    print(f"Created test DOCX file: {output_path}")

if __name__ == "__main__":
    create_test_docx()