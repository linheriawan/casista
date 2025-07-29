#!/usr/bin/env python3
"""
RAG (Retrieval-Augmented Generation) System for Coder Assistant
Provides contextual knowledge retrieval for better coding assistance
"""

import os
import json
import hashlib
from pathlib import Path
from typing import List, Dict, Tuple, Optional
from dataclasses import dataclass
import sqlite3
from rich.console import Console

console = Console()

# Try to import embedding libraries
try:
    import numpy as np
    from sentence_transformers import SentenceTransformer
    HAS_EMBEDDINGS = True
except ImportError:
    HAS_EMBEDDINGS = False
    console.print("[yellow]⚠️ RAG dependencies not installed[/]")
    console.print("[dim]Install with: pip install sentence-transformers numpy faiss-cpu[/]")

# Try to import document processing libraries
try:
    from docx import Document as DocxDocument
    import PyPDF2
    HAS_DOCUMENT_SUPPORT = True
except ImportError:
    HAS_DOCUMENT_SUPPORT = False
    console.print("[yellow]⚠️ Document processing dependencies not installed[/]")
    console.print("[dim]Install with: pip install python-docx PyPDF2[/]")

@dataclass
class Document:
    """Represents a document in the knowledge base"""
    id: str
    content: str
    source: str
    metadata: Dict
    embedding: Optional[List[float]] = None

class RAGKnowledgeBase:
    """Manages the RAG knowledge base with embeddings and retrieval"""
    
    def __init__(self, kb_dir: Path = None):
        self.kb_dir = kb_dir or Path(".ai_context/knowledge")
        self.kb_dir.mkdir(parents=True, exist_ok=True)
        
        self.db_path = self.kb_dir / "knowledge.db"
        self.embeddings_cache = self.kb_dir / "embeddings.json"
        
        # Initialize embedding model
        if HAS_EMBEDDINGS:
            try:
                console.print("[cyan]📥 Loading embedding model...[/]")
                self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')  # Lightweight, good quality
                console.print("[green]✅ Embedding model loaded[/]")
            except Exception as e:
                console.print(f"[red]❌ Failed to load embedding model: {e}[/]")
                self.embedding_model = None
        else:
            self.embedding_model = None
        
        # Initialize database
        self._init_database()
    
    def _init_database(self):
        """Initialize SQLite database for documents"""
        with sqlite3.connect(self.db_path) as conn:
            conn.execute("""
                CREATE TABLE IF NOT EXISTS documents (
                    id TEXT PRIMARY KEY,
                    content TEXT NOT NULL,
                    source TEXT NOT NULL,
                    metadata TEXT,
                    embedding TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """)
            conn.execute("""
                CREATE INDEX IF NOT EXISTS idx_source ON documents(source)
            """)
    
    def add_document(self, content: str, source: str, metadata: Dict = None) -> str:
        """Add a document to the knowledge base"""
        if not content.strip():
            return None
            
        # Generate document ID
        doc_id = hashlib.md5(f"{source}:{content[:100]}".encode()).hexdigest()
        
        # Generate embedding
        embedding = None
        if self.embedding_model:
            try:
                embedding_vector = self.embedding_model.encode(content)
                embedding = embedding_vector.tolist()
            except Exception as e:
                console.print(f"[yellow]⚠️ Failed to generate embedding: {e}[/]")
        
        # Store in database
        with sqlite3.connect(self.db_path) as conn:
            conn.execute("""
                INSERT OR REPLACE INTO documents 
                (id, content, source, metadata, embedding)
                VALUES (?, ?, ?, ?, ?)
            """, (
                doc_id,
                content,
                source,
                json.dumps(metadata or {}),
                json.dumps(embedding) if embedding else None
            ))
        
        return doc_id
    
    def _extract_text_from_docx(self, file_path: Path) -> Tuple[str, Dict]:
        """Extract text content from .docx file"""
        if not HAS_DOCUMENT_SUPPORT:
            return "", {"error": "Document support not available"}
        
        try:
            doc = DocxDocument(file_path)
            text_content = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            content = "\n\n".join(text_content)
            
            # Extract metadata
            props = doc.core_properties
            metadata = {
                "document_type": "docx",
                "author": props.author if props.author else "Unknown",
                "title": props.title if props.title else file_path.stem,
                "created": str(props.created) if props.created else None,
                "modified": str(props.modified) if props.modified else None,
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables)
            }
            
            return content, metadata
            
        except Exception as e:
            console.print(f"[yellow]⚠️ Failed to extract from DOCX {file_path}: {e}[/]")
            return "", {"error": str(e)}
    
    def _extract_text_from_pdf(self, file_path: Path) -> Tuple[str, Dict]:
        """Extract text content from .pdf file"""
        if not HAS_DOCUMENT_SUPPORT:
            return "", {"error": "Document support not available"}
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text_content = []
                
                # Extract text from all pages
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(f"--- Page {page_num + 1} ---\n{page_text.strip()}")
                    except Exception as e:
                        console.print(f"[yellow]⚠️ Failed to extract page {page_num + 1} from {file_path}: {e}[/]")
                
                content = "\n\n".join(text_content)
                
                # Extract metadata
                metadata = {
                    "document_type": "pdf",
                    "pages": len(pdf_reader.pages),
                    "title": file_path.stem
                }
                
                # Try to get PDF metadata
                try:
                    if pdf_reader.metadata:
                        pdf_info = pdf_reader.metadata
                        metadata.update({
                            "author": pdf_info.get('/Author', 'Unknown'),
                            "creator": pdf_info.get('/Creator', 'Unknown'),
                            "title": pdf_info.get('/Title', file_path.stem),
                            "subject": pdf_info.get('/Subject', ''),
                            "creation_date": str(pdf_info.get('/CreationDate', '')),
                            "modification_date": str(pdf_info.get('/ModDate', ''))
                        })
                except:
                    pass  # Metadata extraction failed, use defaults
                
                return content, metadata
                
        except Exception as e:
            console.print(f"[yellow]⚠️ Failed to extract from PDF {file_path}: {e}[/]")
            return "", {"error": str(e)}
    
    def _chunk_text(self, text: str, max_chunk_size: int = 1000, overlap: int = 100) -> List[str]:
        """Split large text into overlapping chunks"""
        if len(text) <= max_chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + max_chunk_size
            
            # Try to break at sentence boundaries
            if end < len(text):
                # Look for sentence endings within the last 200 characters
                search_start = max(start + max_chunk_size - 200, start)
                sentence_end = -1
                for i in range(end - 1, search_start - 1, -1):
                    if text[i] in '.!?':
                        sentence_end = i + 1
                        break
                
                if sentence_end > 0:
                    end = sentence_end
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # Move start position with overlap
            start = end - overlap if end < len(text) else len(text)
        
        return chunks
    
    def index_codebase(self, project_path: Path, extensions: List[str] = None):
        """Index code files in a project"""
        if not project_path.exists():
            console.print(f"[red]❌ Project path not found: {project_path}[/]")
            return
        
        extensions = extensions or ['.py', '.js', '.ts', '.md', '.txt', '.json', '.yaml', '.yml', '.docx', '.pdf']
        
        console.print(f"[cyan]📚 Indexing codebase: {project_path}[/]")
        
        indexed_count = 0
        for file_path in project_path.rglob("*"):
            if (file_path.is_file() and 
                file_path.suffix.lower() in extensions and
                not any(part.startswith('.') for part in file_path.parts[1:]) and  # Skip hidden dirs
                'node_modules' not in str(file_path) and
                '__pycache__' not in str(file_path)):
                
                try:
                    # Handle different file types
                    file_ext = file_path.suffix.lower()
                    
                    if file_ext == '.docx':
                        content, doc_metadata = self._extract_text_from_docx(file_path)
                    elif file_ext == '.pdf':
                        content, doc_metadata = self._extract_text_from_pdf(file_path)
                    else:
                        # Regular text file
                        content = file_path.read_text(encoding='utf-8')
                        doc_metadata = {
                            'file_type': file_ext,
                            'size': len(content)
                        }
                    
                    # Skip empty content or errors
                    if not content or not content.strip():
                        continue
                    
                    # Skip very large content (adjust limit for documents)
                    size_limit = 200000 if file_ext in ['.pdf', '.docx'] else 50000
                    if len(content) > size_limit:
                        console.print(f"[yellow]⚠️ Skipping large file: {file_path.relative_to(project_path)} ({len(content)} chars)[/]")
                        continue
                    
                    # Add common metadata
                    metadata = {
                        **doc_metadata,
                        'relative_path': str(file_path.relative_to(project_path)),
                        'content_size': len(content)
                    }
                    
                    # For large documents, split into chunks
                    if len(content) > 2000 and file_ext in ['.pdf', '.docx']:
                        chunks = self._chunk_text(content, max_chunk_size=2000, overlap=200)
                        for i, chunk in enumerate(chunks):
                            chunk_metadata = {
                                **metadata,
                                'chunk_index': i,
                                'total_chunks': len(chunks),
                                'is_chunked': True
                            }
                            
                            chunk_source = f"{file_path.relative_to(project_path)}#chunk{i}"
                            doc_id = self.add_document(
                                content=chunk,
                                source=chunk_source,
                                metadata=chunk_metadata
                            )
                            
                            if doc_id:
                                indexed_count += 1
                    else:
                        # Index as single document
                        doc_id = self.add_document(
                            content=content,
                            source=str(file_path.relative_to(project_path)),
                            metadata=metadata
                        )
                    
                    if doc_id:
                        indexed_count += 1
                        if indexed_count % 10 == 0:
                            console.print(f"[dim]Indexed {indexed_count} files...[/]")
                
                except Exception as e:
                    console.print(f"[yellow]⚠️ Failed to index {file_path}: {e}[/]")
        
        console.print(f"[green]✅ Indexed {indexed_count} files from codebase[/]")
    
    def retrieve_similar(self, query: str, top_k: int = 5) -> List[Document]:
        """Retrieve most similar documents to query"""
        if not self.embedding_model:
            return self._fallback_keyword_search(query, top_k)
        
        try:
            # Generate query embedding
            query_embedding = self.embedding_model.encode(query)
            
            # Get all documents with embeddings
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.execute("""
                    SELECT id, content, source, metadata, embedding 
                    FROM documents 
                    WHERE embedding IS NOT NULL
                """)
                
                documents = []
                similarities = []
                
                for row in cursor.fetchall():
                    doc_id, content, source, metadata_str, embedding_str = row
                    
                    if embedding_str:
                        doc_embedding = np.array(json.loads(embedding_str))
                        
                        # Calculate cosine similarity
                        similarity = np.dot(query_embedding, doc_embedding) / (
                            np.linalg.norm(query_embedding) * np.linalg.norm(doc_embedding)
                        )
                        
                        documents.append(Document(
                            id=doc_id,
                            content=content,
                            source=source,
                            metadata=json.loads(metadata_str),
                            embedding=doc_embedding.tolist()
                        ))
                        similarities.append(similarity)
                
                # Sort by similarity and return top_k
                if documents:
                    sorted_docs = sorted(zip(documents, similarities), 
                                       key=lambda x: x[1], reverse=True)
                    return [doc for doc, sim in sorted_docs[:top_k]]
                
        except Exception as e:
            console.print(f"[yellow]⚠️ Embedding search failed: {e}[/]")
            return self._fallback_keyword_search(query, top_k)
        
        return []
    
    def _fallback_keyword_search(self, query: str, top_k: int) -> List[Document]:
        """Fallback to simple keyword search when embeddings unavailable"""
        keywords = query.lower().split()
        
        with sqlite3.connect(self.db_path) as conn:
            # Simple keyword matching
            where_conditions = []
            params = []
            
            for keyword in keywords[:3]:  # Limit to first 3 keywords
                where_conditions.append("LOWER(content) LIKE ?")
                params.append(f"%{keyword}%")
            
            if where_conditions:
                query_sql = f"""
                    SELECT id, content, source, metadata
                    FROM documents 
                    WHERE {' OR '.join(where_conditions)}
                    LIMIT ?
                """
                params.append(top_k)
                
                cursor = conn.execute(query_sql, params)
                documents = []
                
                for row in cursor.fetchall():
                    doc_id, content, source, metadata_str = row
                    documents.append(Document(
                        id=doc_id,
                        content=content,
                        source=source,
                        metadata=json.loads(metadata_str)
                    ))
                
                return documents
        
        return []
    
    def get_stats(self) -> Dict:
        """Get knowledge base statistics"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.execute("SELECT COUNT(*) FROM documents")
            total_docs = cursor.fetchone()[0]
            
            cursor = conn.execute("SELECT COUNT(*) FROM documents WHERE embedding IS NOT NULL")
            embedded_docs = cursor.fetchone()[0]
            
            cursor = conn.execute("SELECT DISTINCT source FROM documents")
            sources = [row[0] for row in cursor.fetchall()]
        
        return {
            'total_documents': total_docs,
            'embedded_documents': embedded_docs,
            'unique_sources': len(sources),
            'sources': sources[:10],  # First 10 sources
            'embedding_model': 'all-MiniLM-L6-v2' if HAS_EMBEDDINGS else 'None'
        }

class RAGAssistant:
    """RAG-enhanced assistant that retrieves relevant context"""
    
    def __init__(self, knowledge_base: RAGKnowledgeBase):
        self.kb = knowledge_base
    
    def enhance_prompt(self, user_query: str, system_prompt: str = "") -> str:
        """Enhance prompt with retrieved context"""
        # Retrieve relevant documents
        relevant_docs = self.kb.retrieve_similar(user_query, top_k=3)
        
        if not relevant_docs:
            return system_prompt + f"\n\nUser: {user_query}"
        
        # Build context section
        context_parts = []
        for i, doc in enumerate(relevant_docs, 1):
            source = doc.source
            content = doc.content[:500] + "..." if len(doc.content) > 500 else doc.content
            context_parts.append(f"[Context {i} - {source}]\n{content}")
        
        context_section = "\n\n".join(context_parts)
        
        # Build enhanced prompt
        enhanced_prompt = f"""{system_prompt}

## Relevant Context:
{context_section}

## User Query:
{user_query}

Please provide a helpful response using the above context when relevant. If you reference specific code or information, mention the source."""
        
        return enhanced_prompt
    
    def get_context_summary(self, user_query: str) -> str:
        """Get a summary of retrieved context"""
        relevant_docs = self.kb.retrieve_similar(user_query, top_k=3)
        
        if not relevant_docs:
            return "No relevant context found in knowledge base."
        
        sources = [doc.source for doc in relevant_docs]
        return f"Found {len(relevant_docs)} relevant documents: {', '.join(sources)}"